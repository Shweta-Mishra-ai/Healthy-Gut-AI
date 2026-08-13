import csv
import io
import logging
import os
import re
import zipfile

logger = logging.getLogger("gutfolio.export")

FONTS_DIR = os.path.join(os.path.dirname(__file__), "..", "static", "fonts")
DEVANAGARI_FONT_PATH = os.path.join(FONTS_DIR, "NotoSansDevanagari.ttf")

_INLINE_MD_RE = re.compile(r"(\*\*|__|\*|_|`)")
_TABLE_DIVIDER_RE = re.compile(r"^\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)*\|?$")


def _strip_inline_markdown(text: str) -> str:
    """Removes bold/italic/code markers. Links keep their label, not the URL —
    a raw URL in the middle of a printed sentence is unreadable."""
    text = re.sub(r"\[([^\]]+)\]\(([^)]*)\)", r"\1", text)
    return _INLINE_MD_RE.sub("", text).strip()


def _split_table_row(line: str) -> list[str]:
    cells = line.strip().strip("|").split("|")
    return [_strip_inline_markdown(c.strip()) for c in cells]


def _parse_blocks(markdown_text: str):
    """Walks the markdown once and yields typed blocks:
    ("heading", level, text) | ("bullet", text) | ("table", rows) | ("para", text)

    Tables get their own block type because both exporters previously threw
    them away: the DOCX writer skipped every line starting with '|' outright,
    and the PDF writer stripped the pipes so a comparison table came out as a
    run of jammed-together words. The foods-to-eat / foods-to-avoid table is
    often the single most useful part of one of these articles, so losing it
    silently in an export was real data loss.
    """
    lines = (markdown_text or "").splitlines()
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("|") and stripped.count("|") >= 2:
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = lines[i].strip()
                if not _TABLE_DIVIDER_RE.match(row.strip("|").strip()) and not _TABLE_DIVIDER_RE.match(row):
                    rows.append(_split_table_row(row))
                i += 1
            if rows:
                width = max(len(r) for r in rows)
                yield ("table", [r + [""] * (width - len(r)) for r in rows])
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            yield ("heading", len(heading.group(1)), _strip_inline_markdown(heading.group(2)))
            i += 1
            continue

        if re.match(r"^([-*+]|\d+\.)\s+", stripped):
            yield ("bullet", _strip_inline_markdown(re.sub(r"^([-*+]|\d+\.)\s+", "", stripped)))
            i += 1
            continue

        if set(stripped) <= {"-", "*", "_"} and len(stripped) >= 3:
            i += 1  # horizontal rule
            continue

        yield ("para", _strip_inline_markdown(stripped))
        i += 1


def markdown_to_docx_bytes(title: str, markdown_text: str) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(title or "Article", level=1)

    for block in _parse_blocks(markdown_text):
        kind = block[0]
        if kind == "heading":
            _, level, text = block
            doc.add_heading(text, level=min(level, 4))
        elif kind == "bullet":
            doc.add_paragraph(block[1], style="List Bullet")
        elif kind == "table":
            rows = block[1]
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            for r_idx, row in enumerate(rows):
                for c_idx, cell in enumerate(row):
                    table.cell(r_idx, c_idx).text = cell
            if rows:
                for cell in table.rows[0].cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
        else:
            doc.add_paragraph(block[1])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _has_devanagari(text: str) -> bool:
    return any("ऀ" <= ch <= "ॿ" for ch in text)


def _clean_pdf_text(text: str) -> str:
    """Sanitizes unicode text for the built-in Helvetica font, converting common
    extended unicode characters (curly quotes, dashes, bullets) to clean
    latin-1 equivalents. Only used for the latin-1-only fallback path — text
    that goes through the embedded Devanagari font is left untouched, since
    encoding it to latin-1 would silently delete every Hindi character."""
    if not text:
        return ""
    replacements = {
        "“": '"', "”": '"', "‘": "'", "’": "'", "–": "-", "—": "-",
        "…": "...", "•": "-", "™": "(TM)", "®": "(R)", "©": "(C)",
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return text.encode("latin-1", "ignore").decode("latin-1")


class FontUnavailableError(RuntimeError):
    """The embedded Unicode font needed for this article is missing.

    Raised rather than falling through to Helvetica, because that fallback
    encodes to latin-1 and would hand the user a PDF of blank pages with a
    200 OK and no indication anything went wrong.
    """


def markdown_to_pdf_bytes(title: str, markdown_text: str) -> bytes:
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    # Helvetica (a built-in FPDF core font) only supports latin-1, which
    # silently drops every Devanagari character on encode — a Hindi article
    # exported to PDF used to come out as a page of blank lines with zero
    # error. If the article contains Hindi, embed a real Unicode font
    # instead of trying to transliterate or strip it.
    needs_unicode_font = _has_devanagari(title) or _has_devanagari(markdown_text)
    font_available = os.path.exists(DEVANAGARI_FONT_PATH)

    if needs_unicode_font and not font_available:
        logger.error("Devanagari font missing at %s — refusing to emit a silently blank PDF", DEVANAGARI_FONT_PATH)
        raise FontUnavailableError(
            "This article contains Hindi text, but the embedded Devanagari font "
            "(static/fonts/NotoSansDevanagari.ttf) is missing from the deployment. "
            "Export as DOCX or Markdown instead, or restore the font file."
        )

    use_unicode_font = needs_unicode_font and font_available

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    if use_unicode_font:
        pdf.add_font("NotoDevanagari", "", DEVANAGARI_FONT_PATH)
        base_font = "NotoDevanagari"
        def clean(s):
            return s
    else:
        base_font = "Helvetica"
        clean = _clean_pdf_text

    def set_font(size, bold=False):
        # The embedded Devanagari face is registered in the regular style
        # only; asking fpdf for "B" on it raises rather than synthesising.
        pdf.set_font(base_font, "B" if (bold and not use_unicode_font) else "", size)

    set_font(16, bold=True)
    pdf.multi_cell(0, 10, clean(title or "Article"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)

    epw = pdf.w - 2 * pdf.l_margin

    for block in _parse_blocks(markdown_text):
        kind = block[0]
        if kind == "heading":
            _, level, text = block
            pdf.ln(2)
            set_font(14 if level <= 2 else 12, bold=True)
            pdf.multi_cell(0, 8, clean(text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            set_font(11)
        elif kind == "bullet":
            set_font(11)
            pdf.multi_cell(0, 7, clean(f"- {block[1]}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        elif kind == "table":
            rows = block[1]
            col_width = epw / max(1, len(rows[0]))
            for r_idx, row in enumerate(rows):
                set_font(10, bold=(r_idx == 0))
                # Uniform row height keeps cell borders aligned; long cells
                # are wrapped by fpdf inside their own fixed-width box.
                line_height = 6
                heights = [
                    len(pdf.multi_cell(col_width, line_height, clean(cell), dry_run=True, output="LINES"))
                    for cell in row
                ]
                row_height = line_height * max(heights or [1])
                y_start = pdf.get_y()
                if y_start + row_height > pdf.h - pdf.b_margin:
                    pdf.add_page()
                    y_start = pdf.get_y()
                x = pdf.l_margin
                for cell in row:
                    pdf.set_xy(x, y_start)
                    pdf.multi_cell(col_width, line_height, clean(cell), border=1,
                                   new_x=XPos.RIGHT, new_y=YPos.TOP, max_line_height=line_height)
                    x += col_width
                pdf.set_xy(pdf.l_margin, y_start + row_height)
            pdf.ln(3)
            set_font(11)
        else:
            set_font(11)
            text = clean(block[1])
            if text.strip():
                pdf.multi_cell(0, 7, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    out = pdf.output()
    if isinstance(out, str):
        return out.encode("latin-1", "replace")
    return bytes(out)


def _safe_filename(text: str, fallback: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9\-]+", "-", (text or fallback).strip().lower()).strip("-")
    return slug or fallback


def build_batch_zip(items: list[dict]) -> bytes:
    """Bundles a batch of generation results into one ZIP: one .docx per
    successful article, plus a batch_summary.csv covering every item
    (including failures, so the CSV is a complete audit trail of the run)."""
    buf = io.BytesIO()
    csv_buf = io.StringIO()
    writer = csv.writer(csv_buf)
    writer.writerow([
        "Topic", "Keyword", "Geo", "Status", "Provider", "Words", "Readability",
        "KeywordDensity%", "QualityScore", "Compliance", "Error",
    ])

    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        used_names = set()
        for item in items:
            req = item["request"]
            result = item["result"]
            topic, keyword, geo = req.get("topic", ""), req.get("primary_keyword", ""), req.get("geo_target", "")

            if "error" in result:
                writer.writerow([topic, keyword, geo, "FAILED", "", "", "", "", "", "", result["error"]])
                continue

            article_md = result.get("optimized_article_markdown", "")
            word_count = len(article_md.split())
            readability = result.get("metrics", {}).get("readability", {}).get("fleschReadingEase", "")
            density = result.get("metrics", {}).get("keywordDensity", {}).get("keywordDensityPercent", "")
            quality_score = result.get("quality", {}).get("score", "")
            compliance = result.get("compliance", {}).get("risk_level", "")
            writer.writerow([
                topic, keyword, geo, "OK", result.get("provider_used", ""), word_count,
                readability, density, quality_score, compliance, "",
            ])

            base_name = _safe_filename(topic, f"article-{len(used_names) + 1}")
            filename = f"{base_name}.docx"
            n = 2
            while filename in used_names:
                filename = f"{base_name}-{n}.docx"
                n += 1
            used_names.add(filename)

            docx_bytes = markdown_to_docx_bytes(topic, article_md)
            zf.writestr(filename, docx_bytes)

        zf.writestr("batch_summary.csv", csv_buf.getvalue())

    return buf.getvalue()
